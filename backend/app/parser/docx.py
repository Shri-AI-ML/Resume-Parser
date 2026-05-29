import docx
import logging
from pathlib import Path
from app.parser.base import BaseParser
from app.parser.exceptions import ParserError

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    """
    Parser implementation for DOCX resumes using python-docx.
    """
    def parse(self, file_path: Path) -> tuple[str, int]:
        logger.info(f"File detected: {file_path}")
        
        if not file_path.exists():
            raise ParserError(f"DOCX file does not exist at: {file_path}")

        try:
            # Selected DOCX parser
            logger.info("Parser selected: DocxParser")
            
            # Open document
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraph texts
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Extract table texts
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            raw_text = "\n".join(text_parts).strip()
            
            # Estimate pages: DOCX does not natively store page count.
            # 1 page ≈ 3000 characters (including spaces) is a standard estimation for resumes.
            page_count = max(1, len(raw_text) // 3000)
            
            logger.info(f"Extraction success: {file_path} (Estimated Pages: {page_count}, Bytes: {len(raw_text)})")
            return raw_text, page_count
            
        except Exception as e:
            logger.error(f"Extraction failure: Failed to parse DOCX {file_path}. Reason: {str(e)}")
            raise ParserError(f"Failed to parse DOCX document: {str(e)}")
